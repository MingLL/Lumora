#!/usr/bin/env python3
import argparse
import collections
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    props.append(fill)


def set_cell_text(cell, value, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(value or ""))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=(255, 255, 255))
        shade(table.rows[0].cells[idx], "9C0006")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in (("Title", 24, (31, 78, 121)), ("Heading 1", 16, (31, 78, 121)), ("Heading 2", 13, (156, 0, 6))):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("attack_xlsx", type=Path)
    parser.add_argument("full_xlsx", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    attack_wb = load_workbook(args.attack_xlsx, read_only=True, data_only=True)
    detail = attack_wb["攻击请求明细"]
    headers = [cell.value for cell in next(detail.iter_rows(min_row=1, max_row=1))]
    idx = {name: i for i, name in enumerate(headers)}
    rows = list(detail.iter_rows(min_row=2, values_only=True))

    statuses = collections.Counter(row[idx["状态码"]] for row in rows)
    ips = collections.Counter(row[idx["IP"]] for row in rows)
    categories = collections.Counter()
    for row in rows:
        for category in str(row[idx["攻击分类"]]).split("；"):
            categories[category] += 1

    full_wb = load_workbook(args.full_xlsx, read_only=True, data_only=True)
    ip_sheet = full_wb["IP汇总"]
    ip_headers = [cell.value for cell in next(ip_sheet.iter_rows(min_row=1, max_row=1))]
    ip_idx = {name: i for i, name in enumerate(ip_headers)}
    geo = {}
    for row in ip_sheet.iter_rows(min_row=2, values_only=True):
        geo[row[ip_idx["IP"]]] = row[ip_idx["地点"]]

    summary_sheet = attack_wb["攻击链接汇总"]
    summary_headers = [cell.value for cell in next(summary_sheet.iter_rows(min_row=1, max_row=1))]
    summary_idx = {name: i for i, name in enumerate(summary_headers)}
    link_rows = list(summary_sheet.iter_rows(min_row=2, values_only=True))

    doc = Document()
    configure_doc(doc)
    title = doc.add_heading("dev1 Web 访问安全分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("分析范围：2026-07-27 至 2026-08-14（服务器现存日志约 19 天）")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    doc.add_heading("1. 执行摘要", level=1)
    p = doc.add_paragraph()
    p.add_run("结论：").bold = True
    p.add_run("日志中存在大量自动化漏洞扫描和恶意载荷探测，但筛出的攻击请求没有返回 2xx 成功状态。")
    p = doc.add_paragraph()
    p.add_run("风险判断：").bold = True
    p.add_run("当前证据未显示攻击成功。3,496 次返回 404，27 次返回 400，10 次返回 308 重定向；308 仅表示跳转，不代表命令执行。此前针对 Mozi 的请求也由 Traefik 直接以 404 拒绝，未进入后端。")

    add_table(doc, ["指标", "数量"], [
        ("原始请求总数", "25,326"),
        ("独立访问 IP", "2,455"),
        ("疑似攻击请求", f"{len(rows):,}"),
        ("独立疑似攻击链接", f"{len(link_rows):,}"),
        ("攻击请求占全部请求", f"{len(rows) / 25326:.1%}"),
        ("2xx 攻击响应", str(sum(count for code, count in statuses.items() if 200 <= int(code) < 300))),
    ], widths=[8, 5])

    doc.add_heading("2. 攻击类型分布", level=1)
    doc.add_paragraph("同一请求可能命中多个规则，因此分类数量相加可能大于疑似攻击请求总数。高置信度包含明确载荷；中置信度多为通用扫描，仍需结合响应码判断。")
    category_rows = [(name, f"{count:,}", f"{count / len(rows):.1%}") for name, count in categories.most_common()]
    add_table(doc, ["攻击分类", "请求次数", "占疑似攻击请求比例"], category_rows, widths=[9, 4, 5])

    doc.add_heading("3. 高频攻击链接", level=1)
    doc.add_paragraph("下表列出请求次数最高的 25 个链接。链接仅用于事件复核，请勿在浏览器或服务器终端中直接执行其中参数。")
    top_links = []
    for row in link_rows[:25]:
        url = str(row[summary_idx["链接"]] or "")
        if len(url) > 120:
            url = url[:117] + "…"
        top_links.append((row[summary_idx["攻击分类"]], row[summary_idx["置信度"]],
                          row[summary_idx["请求次数"]], row[summary_idx["状态码分布"]], url))
    add_table(doc, ["分类", "置信度", "次数", "状态码", "链接"], top_links, widths=[4.2, 1.8, 1.6, 2.2, 9.2])

    doc.add_heading("4. 高频来源 IP", level=1)
    source_rows = [(ip, f"{count:,}", geo.get(ip, "")) for ip, count in ips.most_common(20)]
    add_table(doc, ["来源 IP", "疑似攻击请求", "近似地点"], source_rows, widths=[5, 4, 9])

    doc.add_heading("5. 主要攻击行为说明", level=1)
    explanations = [
        ("敏感文件与凭据探测", "尝试访问 .env、.git/config、配置文件、云凭据或应用诊断端点，目的是获取密钥、数据库密码或源代码信息。"),
        ("CMS 漏洞扫描", "集中探测 wp-login.php、xmlrpc.php、WordPress 插件目录等。当前站点并非 WordPress，这类请求主要是互联网批量扫描。"),
        ("WebShell/RCE 扫描", "探测 setup.cgi、phpunit、cgi-bin 及常见 WebShell 路径，尝试远程执行系统命令。"),
        ("命令注入与恶意下载", "URL 参数中包含 wget、shell、系统命令或木马下载地址。Mozi 请求属于此类，但日志显示返回 404 且未转发后端。"),
        ("路径穿越与文件读取", "使用 ../、编码后的目录跳转或 /proc、/etc/passwd 等路径，试图读取服务器文件。"),
        ("备份与源码泄露探测", "扫描 .bak、.sql、.zip、旧配置等可能遗留的备份文件。"),
    ]
    for name, explanation in explanations:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}：").bold = True
        p.add_run(explanation)

    doc.add_heading("6. 风险判断与证据", level=1)
    evidence = [
        "疑似攻击请求全部为 404、400 或 308，没有观察到 2xx 成功响应。",
        "Mozi/Netgear 命令注入请求的 Traefik OriginStatus 为 0，表示没有后端服务处理该请求。",
        "dev1、dev2 主机及运行中容器未发现 Mozi/netgear 文件、异常进程、相关网络连接或持久化项。",
        "Lumora 后端日志中没有命中 setup.cgi、Mozi.m 或相关恶意下载 IP。",
        "当前结论基于服务器保留的约 19 天日志；更早记录因日志轮转不可见。",
    ]
    for item in evidence:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. 建议", level=1)
    recommendations = [
        "保持 Traefik Host 和 Path 路由限制，避免添加不带 Host 的通配入口。",
        "对 .env、.git、备份文件、cgi-bin、setup.cgi 等路径在入口层直接拒绝，并记录来源 IP。",
        "配置访问日志至少保留 30–90 天，最好发送到独立日志系统，避免节点重启或轮转造成证据缺失。",
        "针对短时间大量 404 的 IP 配置速率限制；封禁应基于行为而非永久依赖单个 IP，因为扫描源会频繁变化。",
        "持续监控 2xx/5xx 的高风险路径。一旦高风险载荷返回 2xx，应立即升级为入侵事件排查。",
        "定期检查容器镜像、K3s、Traefik 和操作系统安全更新。",
    ]
    for item in recommendations:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("8. 附件与统计口径", level=1)
    doc.add_paragraph(f"完整访问报告：{args.full_xlsx.name}")
    doc.add_paragraph(f"攻击链接筛选明细：{args.attack_xlsx.name}")
    doc.add_paragraph("筛选采用 URL 特征规则。中置信度结果可能包含无害探测或与业务重名的路径；本报告不把单次扫描等同于攻击成功。IP 地点是网络出口的近似位置，不代表访问者真实所在地。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Lumora · dev1 Web Security Review · 2026-08-14")
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
